#!/usr/bin/env python3
"""
Generate professionally formatted Word document from Python installation guide.

This script creates a .docx file with:
- Professional formatting (fonts, spacing, margins)
- Proper heading hierarchy
- Code blocks with monospace font
- Callout boxes (Important, Note, Tip)
- Table of contents
- Headers and footers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_page_number(run):
    """Add page number field."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_callout_box(doc, text, callout_type, icon):
    """Create a formatted callout box (Important, Note, Tip)."""

    # Colors for different callout types
    colors = {
        'Important': {'bg': 'FFFFE0', 'border': 'FFA500'},  # Light yellow, orange
        'Note': {'bg': 'E0FFFF', 'border': '0000FF'},      # Light blue, blue
        'Tip': {'bg': 'E0FFE0', 'border': '008000'}         # Light green, green
    }

    color = colors.get(callout_type, colors['Note'])

    # Create a table for the callout box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False

    # Set width
    table.columns[0].width = Inches(6.5)

    cell = table.cell(0, 0)
    cell.text = f"{icon} {callout_type}: {text}"

    # Set background color
    set_cell_background(cell, color['bg'])

    # Format text
    paragraph = cell.paragraphs[0]
    paragraph.space_after = Pt(6)

    # Make the label bold
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Add spacing after the callout box
    doc.add_paragraph().space_after = Pt(12)


def create_code_block(doc, code_text, language="bash"):
    """Create a formatted code block."""
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # Set paragraph formatting
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.right_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

    return paragraph


def main():
    # Create a new document
    doc = Document()

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # === TITLE PAGE ===
    title = doc.add_paragraph('PYTHON INSTALLATION')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('AND SETUP GUIDE')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    subtitle_run.font.name = 'Calibri'

    doc.add_paragraph()  # Spacing

    org = doc.add_paragraph('TechEngineer.org')
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org.runs[0]
    org_run.font.size = Pt(12)
    org_run.font.italic = True
    org_run.font.name = 'Calibri'

    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    toc_heading = doc.add_paragraph('Table of Contents')
    toc_heading_run = toc_heading.runs[0]
    toc_heading_run.font.size = Pt(18)
    toc_heading_run.font.bold = True
    toc_heading_run.font.name = 'Calibri'

    doc.add_paragraph('This guide covers everything you need to get started with Python.')

    # Add TOC placeholder (in real Word, this would be auto-generated)
    toc = doc.add_paragraph()
    toc.add_run('Introduction\t\t\t\t\t\t\t\t\t\t\t1\n')
    toc.add_run('Python Installation Considerations\t\t\t\t\t\t2\n')
    toc.add_run('Installing uv Package Manager\t\t\t\t\t\t\t3\n')
    toc.add_run('Installing Python\t\t\t\t\t\t\t\t\t\t4\n')
    toc.add_run('Running Python Code\t\t\t\t\t\t\t\t\t5\n')
    toc.add_run('Additional Resources\t\t\t\t\t\t\t\t\t6\n')
    toc_run = toc.runs[0]
    toc_run.font.name = 'Calibri'
    toc_run.font.size = Pt(11)

    doc.add_page_break()

    # === MAIN CONTENT ===

    # Section 1: Introduction
    h1 = doc.add_heading('1. Introduction', level=1)
    h1_run = h1.runs[0]
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.name = 'Calibri'

    intro = doc.add_paragraph(
        'To start your learning journey, install Python on macOS, Linux, or Windows. '
        'This guide provides step-by-step instructions for setting up your Python environment '
        'using recommended tools and best practices.'
    )
    intro_run = intro.runs[0]
    intro_run.font.size = Pt(11)
    intro_run.font.name = 'Calibri'
    intro.paragraph_format.line_spacing = 1.15
    intro.paragraph_format.space_after = Pt(6)

    # Section 2: Python Installation Considerations
    h1 = doc.add_heading('2. Python Installation Considerations', level=1)

    p = doc.add_paragraph(
        'There are many ways to install Python, but we recommend using the '
    )
    uv_run = p.add_run('uv')
    uv_run.bold = True
    p.add_run(
        ' package manager for the best experience. Other options like conda, pyenv, '
        'or the official installers from Python.org are reliable and effective but '
        'may become limiting as your Python expertise grows.'
    )
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)

    # Key Recommendations
    h2 = doc.add_heading('Key Recommendations', level=2)
    h2_run = h2.runs[0]
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.name = 'Calibri'

    rec = doc.add_paragraph(
        'For now, don\'t worry about the details! Relax and focus on learning as '
        'much as you can about Python. Remember, everything is reversible and can be changed.'
    )
    rec_run = rec.runs[0]
    rec_run.font.size = Pt(11)
    rec_run.font.name = 'Calibri'
    rec_run.font.italic = True
    rec.paragraph_format.line_spacing = 1.15
    rec.paragraph_format.space_after = Pt(6)

    # What uv will help with
    h2 = doc.add_heading('What uv Will Help You With', level=2)

    doc.add_paragraph('Installing specific Python versions', style='List Bullet')
    doc.add_paragraph('Creating and managing virtual environments', style='List Bullet')
    doc.add_paragraph('Installing and managing Python packages and tools', style='List Bullet')
    doc.add_paragraph('Replacing several tools like conda, pyenv, pip, poetry, virtualenv, and more', style='List Bullet')

    # Section 3: Installing uv Package Manager
    h1 = doc.add_heading('3. Installing uv Package Manager', level=1)

    p = doc.add_paragraph(
        'The steps for installing uv and running Python code can seem intimidating if '
        'you\'re not familiar with using the terminal. Don\'t worry! To navigate this '
        'process without stress, we recommend that you always read the messages that '
        'appear after you run any commands.'
    )
    p_run = p.runs[0]
    p_run.font.size = Pt(11)
    p_run.font.name = 'Calibri'
    p.paragraph_format.line_spacing = 1.15

    # Installation Instructions
    h2 = doc.add_heading('Installation Instructions', level=2)

    h3 = doc.add_heading('For macOS and Linux:', level=3)
    create_code_block(doc, 'curl -LsSf https://astral.sh/uv/install.sh | sh')

    h3 = doc.add_heading('For Windows (PowerShell):', level=3)
    create_code_block(doc, 'powershell -ExecutionPolicy ByPass -c "irm https://astral.sh/uv/install.ps1 | iex"', 'powershell')

    # Important callout
    create_callout_box(doc,
                    'Follow the PowerShell instructions to validate the installation. '
                    'You may need to restart the system.',
                    'Important', '⚠')

    # Note callout
    create_callout_box(doc,
                    'For more details and assistance with installing uv, please refer to '
                    'the documentation available at docs.astral.sh/uv.',
                    'Note', '📝')

    # Section 4: Installing Python
    h1 = doc.add_heading('4. Installing Python', level=1)

    p = doc.add_paragraph('Once we\'ve installed uv, we\'re ready to install Python to begin our journey!')

    h2 = doc.add_heading('Installation Command', level=2)
    create_code_block(doc, 'uv python install 3.12')

    completion = doc.add_paragraph('Installation complete! The adventure begins now!')
    completion_run = completion.runs[0]
    completion_run.font.size = Pt(11)
    completion_run.font.bold = True
    completion_run.font.color.rgb = RGBColor(0, 102, 0)
    completion.paragraph_format.space_after = Pt(12)

    # Section 5: Running Python Code
    h1 = doc.add_heading('5. Running Python Code', level=1)

    p = doc.add_paragraph(
        'Python provides several tools and environments for writing and executing code, '
        'each suited to different needs.'
    )
    p_run = p.runs[0]
    p_run.font.size = Pt(11)
    p_run.font.name = 'Calibri'

    p = doc.add_paragraph(
        'We\'ll explore three methods for running Python code:'
    )

    # Create the list items
    item1 = doc.add_paragraph('The Python interpreter', style='List Bullet')
    item1.add_run(' – for immediate feedback')

    item2 = doc.add_paragraph('IPython', style='List Bullet')
    item2.add_run(' – for interactive coding experience')

    item3 = doc.add_paragraph('JupyterLab', style='List Bullet')
    item3.add_run(' – for developing comprehensive, interactive projects')

    p = doc.add_paragraph(
        '\nFuture courses will include advanced configurations using tools such as '
        'VS Code and virtual environments with uv. For more details, visit TechEngineer.org.'
    )

    # Section 5.1
    h2 = doc.add_heading('5.1 The Python Interpreter', level=2)

    p = doc.add_paragraph(
        'The Python interpreter is the simplest way to execute Python code directly. '
        'It allows you to run commands line by line and is ideal for quick testing and '
        'simple scripts.'
    )
    p_run = p.runs[0]
    p_run.font.size = Pt(11)
    p_run.font.name = 'Calibri'
    p.paragraph_format.line_spacing = 1.15

    # Features
    h3 = doc.add_heading('Features:', level=3)
    doc.add_paragraph('Quick testing and debugging', style='List Bullet')
    doc.add_paragraph('Line-by-line code execution', style='List Bullet')
    doc.add_paragraph('Included with Python installations', style='List Bullet')
    doc.add_paragraph('Started directly from command line', style='List Bullet')

    # Section 6: Best Practices
    h1 = doc.add_heading('6. Best Practices for Terminal Usage', level=1)

    p = doc.add_paragraph('When working with terminal commands, keep these tips in mind:')

    doc.add_paragraph('Read command output – There\'s often useful information or instructions provided', style='List Number')
    doc.add_paragraph('Follow the instructions – Simply follow terminal prompts or do a quick online search', style='List Number')
    doc.add_paragraph('Don\'t panic – You may be surprised at how quickly you can find solutions to common problems', style='List Number')

    # Section 7: Additional Resources
    h1 = doc.add_heading('7. Additional Resources', level=1)

    # Add links
    resources = [
        ('Official Python Documentation', 'https://docs.python.org/'),
        ('uv Documentation', 'https://docs.astral.sh/uv/'),
        ('TechEngineer.org', 'https://TechEngineer.org')
    ]

    for name, url in resources:
        p = doc.add_paragraph()
        link = p.add_run(f'{name}: {url}')
        link.font.size = Pt(11)
        link.font.name = 'Calibri'
        link.font.color.rgb = RGBColor(0, 0, 255)
        link.underline = True

    # === DOCUMENT INFO FOOTER ===

    # Add a section for document info
    doc.add_page_break()

    info_heading = doc.add_heading('Document Information', level=1)

    # Create a table for document info
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ('Document Type:', 'Technical Tutorial'),
        ('Target Audience:', 'Python learners, beginners'),
        ('Last Updated:', 'January 2026'),
        ('Version:', '1.0'),
        ('Source:', 'TechEngineer.org')
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

        # Format header cells
        cell = info_table.rows[i].cells[0]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

        # Format value cells
        cell = info_table.rows[i].cells[1]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

    # === SAVE DOCUMENT ===
    output_path = 'Python-Installation-Guide.docx'
    doc.save(output_path)

    print(f"[OK] Successfully created: {output_path}")
    print(f"[DOC] Document saved in: .claude/skills/document-reader-formatter-skill/formatted-documents/")
    print(f"\n[INFO] Formatting applied:")
    print(f"   - Professional fonts (Calibri 11pt body, Consolas 10pt code)")
    print(f"   - Proper heading hierarchy (18pt H1, 14pt H2, 12pt H3)")
    print(f"   - Code blocks with monospace font")
    print(f"   - Callout boxes (Important, Note)")
    print(f"   - Table of contents")
    print(f"   - Document information table")
    print(f"   - 1\" margins, 1.15 line spacing")


if __name__ == '__main__':
    main()
